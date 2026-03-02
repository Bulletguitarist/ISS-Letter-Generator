"""
generate_templates.py
Generates DOCX templates for:
  1. Offer Letter
  2. Internship Certificate
  3. Letter of Recommendation (LOR)

Letterhead: Logo (left) | Office info (right) | Blue border line
Run: python generate_templates.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_DIR   = "templates"
LOGO_PATH = "logo.png"
os.makedirs(OUT_DIR, exist_ok=True)


# ─── Helpers ──────────────────────────────────────────────────────────────────

def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def remove_cell_borders(cell):
    """Remove all borders from a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_bg(cell, color_hex: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  color_hex)
    tcPr.append(shd)


def add_company_header(doc: Document):
    """
    Letterhead matching the screenshot:
      [LOGO]  | Office: Pune, Maharashtra, India | info@innovativestaffingsolutions.online
      _____________________________________________________________________ (blue line)
    """
    # ── 2-column borderless table ─────────────────────────────────────────────
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set table borders to none via tblPr — insert AFTER tblW and jc
    tbl_elem = tbl._tbl
    tblPr = tbl_elem.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_elem.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    # Find tblLook and insert tblBorders before it (correct schema order)
    tblLook = tblPr.find(qn("w:tblLook"))
    if tblLook is not None:
        tblLook.addprevious(tblBorders)
    else:
        tblPr.append(tblBorders)

    left_cell  = tbl.rows[0].cells[0]
    right_cell = tbl.rows[0].cells[1]

    # Remove cell borders
    for cell in [left_cell, right_cell]:
        remove_cell_borders(cell)

    # Set column widths via tcW (already in tcPr from python-docx default)
    # Override left cell width
    left_tcPr = left_cell._tc.get_or_add_tcPr()
    tcW_left  = left_tcPr.find(qn("w:tcW"))
    if tcW_left is not None:
        tcW_left.set(qn("w:w"),    "1584")
        tcW_left.set(qn("w:type"), "dxa")
    right_tcPr = right_cell._tc.get_or_add_tcPr()
    tcW_right  = right_tcPr.find(qn("w:tcW"))
    if tcW_right is not None:
        tcW_right.set(qn("w:w"),    "7416")
        tcW_right.set(qn("w:type"), "dxa")

    # ── Left cell: Logo ───────────────────────────────────────────────────────
    logo_para = left_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_para.paragraph_format.space_after  = Pt(0)
    logo_para.paragraph_format.space_before = Pt(0)

    if os.path.exists(LOGO_PATH):
        run = logo_para.add_run()
        run.add_picture(LOGO_PATH, width=Inches(0.9))
    else:
        r = logo_para.add_run("ISS")
        r.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

    # ── Right cell: Office info ───────────────────────────────────────────────
    right_cell.paragraphs[0].clear()
    tc   = right_cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    tcPr.append(vAlign)

    info_para = right_cell.paragraphs[0]
    info_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    info_para.paragraph_format.space_after  = Pt(0)
    info_para.paragraph_format.space_before = Pt(0)
    r = info_para.add_run(
        "| Office: Pune, Maharashtra, India  |  info@innovativestaffingsolutions.online"
    )
    r.font.size      = Pt(9)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ── Blue horizontal rule — pBdr BEFORE spacing in pPr ────────────────────
    rule = doc.add_paragraph()
    p    = rule._p
    pPr  = p.get_or_add_pPr()

    # pBdr must come before w:spacing in schema order
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "12")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1a3a5c")
    pBdr.append(bot)
    pPr.insert(0, pBdr)   # insert at beginning so ordering is correct


def _make_tcW(twips: int):
    """Helper: create w:tcW element with fixed width in twips."""
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(twips))
    tcW.set(qn("w:type"), "dxa")
    return tcW


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)


def add_body_para(doc: Document, text: str):
    p   = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def add_footer(doc: Document):
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.add_run("For Innovative Staffing Solutions").bold = True
    sig.runs[0].font.size = Pt(10.5)

    doc.add_paragraph()
    doc.add_paragraph()   # signature space

    auth = doc.add_paragraph()
    auth.add_run("Authorized Signatory").bold = True
    auth.runs[0].font.size = Pt(10)

    for label, value in [
        ("Name",        "Kaif Khan"),
        ("Designation", "Founder & Proprietor"),
        ("Date",        "{{today_date}}"),
    ]:
        p  = doc.add_paragraph()
        rb = p.add_run(f"{label}: ")
        rb.bold = True
        rb.font.size = Pt(10)
        p.add_run(value).font.size = Pt(10)

    doc.add_paragraph()
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = contact.add_run(
        "hr@innovativestaffingsolutions.online  |  +91 7447802076  |  innovativestaffingsolutions.online"
    )
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


# ─── 1. Offer Letter ──────────────────────────────────────────────────────────

def make_offer_letter():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    add_company_header(doc)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("OFFER LETTER \u2013 {{designation}}")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = date_p.add_run("Date: {{today_date}}")
    r2.bold = True
    r2.font.size = Pt(11)

    doc.add_paragraph()

    # Candidate info
    for label, placeholder in [
        ("Candidate Name", "{{name}}"),
        ("Email",          "{{email}}"),
        ("Phone",          "{{phone}}"),
        ("Address",        "{{address}}"),
    ]:
        p  = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        rl = p.add_run(f"{label}: ")
        rl.bold = True
        rl.font.size = Pt(10.5)
        p.add_run(placeholder).font.size = Pt(10.5)

    doc.add_paragraph()

    # Opening
    opening = doc.add_paragraph()
    opening.add_run("We are pleased to offer you the position of ").font.size = Pt(10.5)
    r = opening.add_run("{{designation}}")
    r.bold = True; r.font.size = Pt(10.5)
    opening.add_run(" at ").font.size = Pt(10.5)
    r2 = opening.add_run("Innovative Staffing Solutions")
    r2.bold = True; r2.font.size = Pt(10.5)
    opening.add_run(", effective ").font.size = Pt(10.5)
    r3 = opening.add_run("{{joining_date}}")
    r3.bold = True; r3.font.size = Pt(10.5)
    opening.add_run(".").font.size = Pt(10.5)

    # AI content
    ai_p = doc.add_paragraph()
    ai_p.add_run("{{ai_generated_content}}").font.size = Pt(10.5)

    # Responsibilities
    resp = doc.add_paragraph()
    resp.add_run("During this internship, your responsibilities will include:").font.size = Pt(10.5)

    for item in [
        "Assisting in designing and developing projects in the {{domain}} domain",
        "Collaborating with the team on real-time tasks and deliverables",
        "Participating in reviews, documentation, and team discussions",
        "Contributing to research, analysis, and process improvements",
        "Supporting senior team members in project execution",
        "Maintaining daily reports and updating progress regularly",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph()

    add_body_para(doc,
        "This internship position is project-based and remotely coordinated unless stated otherwise. "
        "Your performance and learning progress will be evaluated throughout the internship, and "
        "high-performing interns may receive extended opportunities or full-time consideration."
    )
    add_body_para(doc,
        "We are excited to welcome you to Innovative Staffing Solutions and look forward to your "
        "meaningful contributions."
    )

    add_footer(doc)

    path = os.path.join(OUT_DIR, "Offer_Letter_Template.docx")
    doc.save(path)
    print(f"✅ Saved: {path}")


# ─── 2. Internship Certificate ────────────────────────────────────────────────

def make_internship_certificate():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    add_company_header(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("INTERNSHIP COMPLETION CERTIFICATE")
    r.bold = True; r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta.add_run("Certificate No.: ISS/INT/{{current_year}}/001").font.size = Pt(9)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.add_run("Date: {{today_date}}").font.size = Pt(9)

    doc.add_paragraph()

    concern = doc.add_paragraph()
    concern.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = concern.add_run("TO WHOM IT MAY CONCERN")
    r.bold = True; r.underline = True; r.font.size = Pt(11)

    doc.add_paragraph()

    # Body
    body1 = doc.add_paragraph()
    body1.add_run("This is to certify that ").font.size = Pt(11)
    r = body1.add_run("{{name}}")
    r.bold = True; r.font.size = Pt(11)
    body1.add_run(
        " has successfully completed an internship with Innovative Staffing Solutions "
        "in the domain of "
    ).font.size = Pt(11)
    r2 = body1.add_run("{{domain}}")
    r2.bold = True; r2.font.size = Pt(11)
    body1.add_run(" from ").font.size = Pt(11)
    r3 = body1.add_run("{{joining_date}}")
    r3.bold = True; r3.font.size = Pt(11)
    body1.add_run(" to ").font.size = Pt(11)
    r4 = body1.add_run("{{last_working_date}}")
    r4.bold = True; r4.font.size = Pt(11)
    body1.add_run(", with a total duration of ").font.size = Pt(11)
    body1.add_run("{{duration}}").font.size = Pt(11)
    body1.add_run(".").font.size = Pt(11)

    doc.add_paragraph()

    body2 = doc.add_paragraph()
    rb = body2.add_run("{{name}}")
    rb.bold = True; rb.font.size = Pt(11)
    body2.add_run(" {{ai_internship_description}}").font.size = Pt(11)

    doc.add_paragraph()

    perf = doc.add_paragraph()
    perf.add_run("Performance: ").bold = True
    perf.runs[0].font.size = Pt(11)
    perf.add_run("{{ai_performance_summary}}").font.size = Pt(11)

    doc.add_paragraph()

    skills_p = doc.add_paragraph()
    skills_p.add_run("Skills Demonstrated: ").bold = True
    skills_p.runs[0].font.size = Pt(11)
    skills_p.add_run("{{ai_skills_summary}}").font.size = Pt(11)

    doc.add_paragraph()

    close = doc.add_paragraph()
    close.add_run(
        "We wish {{name}} all the best in their future endeavors and recommend them "
        "for any future professional assignments."
    ).font.size = Pt(11)

    add_footer(doc)

    path = os.path.join(OUT_DIR, "Internship_Certificate_Template.docx")
    doc.save(path)
    print(f"✅ Saved: {path}")


# ─── 3. LOR ───────────────────────────────────────────────────────────────────

def make_lor():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    add_company_header(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("LETTER OF RECOMMENDATION")
    r.bold = True; r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

    doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.add_run("Date: ").bold = True
    date_p.runs[0].font.size = Pt(10.5)
    date_p.add_run("{{today_date}}").font.size = Pt(10.5)

    doc.add_paragraph()

    concern = doc.add_paragraph()
    concern.add_run("To Whomsoever It May Concern,").bold = True
    concern.runs[0].font.size = Pt(11)

    doc.add_paragraph()

    subj = doc.add_paragraph()
    subj.add_run("Subject: Letter of Recommendation for ").bold = True
    subj.runs[0].font.size = Pt(11)
    r = subj.add_run("{{name}}")
    r.bold = True; r.font.size = Pt(11)

    doc.add_paragraph()

    ai_p = doc.add_paragraph()
    ai_p.add_run("{{ai_lor_body}}").font.size = Pt(11)

    doc.add_paragraph()

    close = doc.add_paragraph()
    close.add_run(
        "We recommend {{name}} without reservation and are confident they will be "
        "a valuable addition to any organization they choose to join."
    ).font.size = Pt(11)

    doc.add_paragraph()

    # Candidate details table
    tbl = doc.add_table(rows=6, cols=2)
    tbl.style = "Table Grid"
    details = [
        ("Full Name",   "{{name}}"),
        ("Designation", "{{designation}}"),
        ("Domain",      "{{domain}}"),
        ("Period",      "{{joining_date}} to {{last_working_date}}"),
        ("Performance", "{{performance}}"),
        ("Email",       "{{email}}"),
    ]
    for i, (label, value) in enumerate(details):
        row = tbl.rows[i].cells
        row[0].text = label
        row[1].text = value
        row[0].paragraphs[0].runs[0].bold = True
        row[0].paragraphs[0].runs[0].font.size = Pt(10)
        row[1].paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(row[0], "EBF3FB")

    add_footer(doc)

    path = os.path.join(OUT_DIR, "LOR_Template.docx")
    doc.save(path)
    print(f"✅ Saved: {path}")


# ─── Main ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    make_offer_letter()
    make_internship_certificate()
    make_lor()
    print("\n✅ All 3 templates generated in /templates folder.")
    print("Run the app:  streamlit run app.py")